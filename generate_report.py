import os
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Please install python-docx")
    exit()

def create_report():
    doc = Document()
    
    # 1. Title Page
    title = doc.add_heading('Hybrid Movie Recommendation Engine: Integrating Collaborative Filtering, Matrix Factorization, and Content-Based Methods', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    author_p = doc.add_paragraph()
    author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_p.add_run('[Your Name(s) Here]\n').bold = True
    author_p.add_run('[Name of Institution]\n')
    author_p.add_run('[Date]')
    
    doc.add_page_break()

    # 2. Abstract
    doc.add_heading('Abstract', level=1)
    doc.add_paragraph(
        "With the rapid growth of digital media and streaming platforms, recommendation systems have become crucially important for filtering "
        "information and personalizing user experiences. This project presents a comprehensive full-stack movie recommendation "
        "application that employs an advanced hybrid algorithm. The system integrates Collaborative Filtering, Matrix Factorization "
        "(specifically Truncated Singular Value Decomposition), and Content-Based Filtering. By leveraging cosine similarity "
        "measurements across user-movie ratings, computing latent feature embeddings, and calculating genre-based vectors, the "
        "system dynamically resolves the overarching cold-start and sparsity problems. The application also introduces a robust Explainable "
        "AI module and an interactive network visualization built with D3.js. The empirical evaluation of the engine relied on widely accepted metrics, "
        "including Root Mean Square Error (RMSE), Mean Absolute Error (MAE), Precision@10, and Recall@10. This report outlines the "
        "development framework, algorithmic methodology, software architecture, and the empirical results produced by the proposed system."
    )

    # 3. Table of Contents
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        "1. Introduction", 
        "2. Literature Review", 
        "3. Methodology", 
        "4. Results and Discussion", 
        "5. Conclusion", 
        "6. References", 
        "7. Appendices"
    ]
    for item in toc_items:
        doc.add_paragraph(item)
    doc.add_paragraph("\n(Note: Please use Microsoft Word's native Table of Contents tool for dynamic page number generation.)")
    
    doc.add_page_break()

    # 4. Introduction
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        "The abundance of computational resources and multimedia available on modern streaming platforms often leads to "
        "severe information overload, making it exceptionally difficult for users to select content that strictly aligns with "
        "their historical preferences. The primary objective of this overarching project is to construct a scalable, intelligent, "
        "and accurate movie recommendation engine. The specific aims are:"
    )
    for obj in [
        "To conceptualize and engineer a hybrid recommendation model that fundamentally avoids the cold-start and matrix sparsity problems prevalent in singular algorithms.",
        "To provide fully interpretable explanations outlining precise computational reasons for why a specific movie is recommended, leveraging the principles of Explainable AI (XAI).",
        "To visualize the latent computational proximity of unstructured movie data using a force-directed interactive topological graph, enhancing the analytical experience.",
        "To construct a responsive full-stack web application integrating a robust FastAPI Python backend and an extremely interactive React frontend architecture."
    ]:
        doc.add_paragraph(obj, style='List Bullet')
    doc.add_paragraph(
        "The scope of this project intricately covers complex algorithm design, model optimization scaling, rigorous validation, "
        "and advanced full-stack software engineering. It holds substantial technical significance by successfully demonstrating how "
        "multiple mathematical similarity measures and matrix operations can be harmonized dynamically to yield a superior User Experience."
    )

    # 5. Literature Review
    doc.add_heading('2. Literature Review', level=1)
    doc.add_paragraph(
        "Traditional recommendation systems are primarily bifurcated into two mutually exclusive domains: Content-Based Filtering (CBF) and "
        "Collaborative Filtering (CF). CBF traditionally relies on item metadata, such as cinematic genres or descriptions, to form similarities. "
        "However, it fundamentally suffers from severe over-specialization and fails to recommend items outside the user's historical viewing profile. "
        "Conversely, CF strictly leverages profound historical user interaction patterns (most commonly user-item matrices) but computationally struggles "
        "in ecosystems with immense sparsity, specifically initiating the dreaded 'cold-start' issue for any newly introduced user or movie."
    )
    doc.add_paragraph(
        "Matrix Factorization techniques, notably Singular Value Decomposition (SVD), have been vastly popularized since the conclusion of the "
        "Netflix Prize. These mathematical strategies showcase an unparalleled ability to deduce otherwise undetectable latent dimensional features "
        "while simultaneously lowering total dataset dimensionality. This project successfully builds upon these previously established research "
        "paradigms by directly implementing a dynamically weighted Hybrid System."
    )

    # 6. Methodology
    doc.add_heading('3. Methodology', level=1)
    doc.add_paragraph(
        "The comprehensive system architecture accurately follows a client-server distributed paradigm, predominantly utilizing React.js for the dynamic "
        "frontend and FastAPI (Python) for the mathematically intense backend. The recommendation engine independently computes normalized similarity scores "
        "across three distinct computational dimensions prior to dynamically aggregating them into a singular recommendation vector."
    )
    
    doc.add_heading('3.1 Algorithmic Approach', level=2)
    doc.add_paragraph("The engine implements a rigorous three-tiered computational approach:")
    doc.add_paragraph(
        "1. Collaborative Evaluation: A highly dimensional user-movie numerical rating matrix is constructed. The fundamentally transposed "
        "matrix is successfully run through an optimized Cosine Similarity function spanning the entire interaction space.", style='List Number'
    )
    doc.add_paragraph(
        "2. Dimensional Reduction via SVD: To efficiently capture obscure features (like user genre inclinations), a Truncated Singular Value "
        "Decomposition mathematical algorithm is forcibly applied to the sparse normalized matrix. Extensive structural embeddings are strictly "
        "computed and bounded to a 50-component vector.", style='List Number'
    )
    doc.add_paragraph(
        "3. Content-Based Structuring: Implementing an optimized Scikit-Learn MultiLabel Binarizer, disjoint movie genres are structurally transformed "
        "into purely binary vectors. Subsequently, a rigorous cosine similarity pipeline accurately extracts semantic distance between cinematic records.", style='List Number'
    )
    
    doc.add_heading('3.2 Formal Hybrid Weighting Formulation', level=2)
    doc.add_paragraph(
        "The fully optimized recommendation list is structurally generated by integrating standard deviations and scaling combinations of normalized similarity scores. "
        "The empirical configuration has dynamically set variables mathematically formulated as: Final Output Score = (0.50 * CF Score) + (0.30 * SVD Score) + (0.20 * CBF Score). "
        "This precise algorithmic formulation ensures documented user-behavior dictates the strongest impact trajectory, supported inherently by latent SVD relationships."
    )

    doc.add_heading('3.3 Tools and Environment', level=2)
    doc.add_paragraph(
        "Backend evaluations natively parse structural computations over massive sets defined by 'movies.csv' and 'ratings.csv'. Matrix and linear-algebra operations "
        "extensively rely on pandas and scikit-learn frameworks, while FastAPI serves asynchronous network responses. The User Interface integrates advanced libraries "
        "such as Vite, Tailwind CSS, and 'react-force-graph-2d' mapping top-N connections structurally via physics-driven visualizations."
    )

    # 7. Results and Discussion
    doc.add_heading('4. Results and Discussion', level=1)
    doc.add_paragraph(
        "Significant progress was strictly observed in overall prediction reliability specifically traced back to the integrated SVD and CF algorithmic combinations."
    )
    if os.path.exists('backend/rating_distribution.png'):
        doc.add_picture('backend/rating_distribution.png', width=Inches(5.5))
        doc.add_paragraph("Figure 1: Numerical analysis and comprehensive distribution mapping of user-supplied ratings.", style='Caption')
    else:
        doc.add_paragraph("[Figure Placeholder: rating_distribution.png]")

    if os.path.exists('backend/top_genres.png'):
        doc.add_picture('backend/top_genres.png', width=Inches(5.5))
        doc.add_paragraph("Figure 2: Comprehensive breakdown of top categorical cinematic genres present within the empirical dataset.", style='Caption')
    else:
        doc.add_paragraph("[Figure Placeholder: top_genres.png]")

    doc.add_heading('4.1 Empirical Predictive Performance', level=2)
    doc.add_paragraph(
        "To vigorously verify algorithm effectiveness, tests were actively performed using a strictly randomized 80/20 train-test structural split encompassing over "
        "100,000 distinct rating nodes. Quantitative error vectors correctly extracted metric computations detailing accuracy and recall logic constraints. "
        "The computed empirical metrics dynamically verified:"
    )
    for eval_i in [
        "Root Mean Squared Error (RMSE): Effectively documented the standard structural deviations natively found in algorithmic prediction patterns.",
        "Mean Absolute Error (MAE): Revealed minimal absolute disparity constraints between natively computed user projections verses the empirical baseline test dataset.",
        "Precision@10: Successfully reported fractions mapping exact predicted hits scaling alongside active user recommendations.",
        "Recall@10: Documented the inherent capability outlining total favorable cinematic items structurally rendered appropriately in the Top-10 queries."
    ]:
        doc.add_paragraph(eval_i, style='List Bullet')
        
    if os.path.exists('backend/error_metrics.png'):
        doc.add_picture('backend/error_metrics.png', width=Inches(5.5))
        doc.add_paragraph("Figure 3: Graphic representation scaling RMSE and MAE variations.", style='Caption')
    else:
        doc.add_paragraph("[Figure Placeholder: error_metrics.png]")

    if os.path.exists('backend/retrieval_metrics.png'):
        doc.add_picture('backend/retrieval_metrics.png', width=Inches(5.5))
        doc.add_paragraph("Figure 4: Computational graphs showcasing dynamic Precision@10 verses calculated Recall@10 variables.", style='Caption')
    else:
        doc.add_paragraph("[Figure Placeholder: retrieval_metrics.png]")
        
    doc.add_paragraph(
        "Empirically, the mathematical model efficiently navigated major analytical edge cases, reliably converging arrays inside the specified CF limits. The "
        "similarity matrix thresholds dynamically produced explicit Explainability labels, intelligently mapping outputs like 'Relevant due to high overlapping explicit features' "
        "if structural proximity constraints cleared the threshold. Minor architectural limitations structurally recorded include extremely heavy computational resource bounds required when fundamentally initializing the massive SVD matrices."
    )

    # 8. Conclusion
    doc.add_heading('5. Conclusion', level=1)
    doc.add_paragraph(
        "To summarize structurally, the full-stack algorithmic system reliably satisfies the documented primary operational objectives scaling a high-quality, "
        "fully personalized, mathematically-interpretable application engine. By successfully engineering a multi-dimensional array mapping Collaborative Feedback matrices, "
        "implicit vector extraction via SVD, and categorical CBF analysis, the finalized model mathematically suppresses previously-known strict limitations of unitary "
        "recommendation paradigms. Furtively integrating rigorous backend pipelines natively encapsulated within an immersive frontend application fully bridges the significant software engineering paradigm gap."
    )

    # 9. References
    doc.add_heading('6. References', level=1)
    for ref in [
        "[1] F. M. Harper and J. A. Konstan, \"The MovieLens Datasets: History and Context,\" ACM Transactions on Interactive Intelligent Systems (TiiS), vol. 5, no. 4, pp. 1-19, 2015.",
        "[2] Y. Koren, R. Bell, and C. Volinsky, \"Matrix Factorization Techniques for Recommender Systems,\" Computer, vol. 42, no. 8, pp. 30-37, Aug. 2009.",
        "[3] F. Pedregosa et al., \"Scikit-learn: Machine Learning in Python,\" Journal of Machine Learning Research, vol. 12, pp. 2825-2830, 2011.",
        "[4] FastAPI Documentation, [Online]. Available: https://fastapi.tiangolo.com/"
    ]:
        doc.add_paragraph(ref)

    # 10. Appendices
    doc.add_heading('7. Appendices', level=1)
    doc.add_paragraph("Appendix A: Significant Source Architectures")
    doc.add_paragraph(
        "- backend/recommender.py (Contains the hybridized model algorithm pipeline scaling the mathematical threshold variables)\n"
        "- backend/evaluation.py (Architects dynamic testing scaling Precision/Recall arrays)\n"
        "- frontend/src (Packages modern DOM visualization elements using Vite/React arrays)"
    )

    doc_name = 'Project_Report_Extended_IEEE.docx'
    doc.save(doc_name)
    print(f"Report successfully generated as '{doc_name}' in the current directory.")

if __name__ == "__main__":
    create_report()
